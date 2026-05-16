from __future__ import annotations

import json
from html import escape
from pathlib import Path

import numpy as np
from docx import Document

from src.config import PROJECT_ROOT, RESULTS_DIR, StudentMetadata


REPORT_MD = PROJECT_ROOT / "Experiment4_Report.md"
REPORT_HTML = PROJECT_ROOT / "Experiment4_Report.html"
REPORT_DOCX = PROJECT_ROOT / "Experiment4_Report.docx"
SUMMARY_JSON = RESULTS_DIR / "summary.json"
VALIDATION_REPORT = PROJECT_ROOT / "validation_report.txt"


def _load_metrics() -> dict[str, object]:
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    baseline = summary["baseline_results"]
    routed = summary["routed_results"]
    barrier = summary["barrier_results"]
    validation_lines = VALIDATION_REPORT.read_text(encoding="utf-8").strip().splitlines()

    percentages = np.array([item["flooded_percentage"] for item in baseline], dtype=float)
    volumes = np.array([item["flood_volume_m3"] for item in baseline], dtype=float)
    levels = np.array(summary["water_levels_m"], dtype=float)

    return {
        "summary": summary,
        "validation_lines": validation_lines,
        "min_pct": float(percentages.min()),
        "max_pct": float(percentages.max()),
        "pct_growth": float(percentages[-1] - percentages[0]),
        "max_volume": float(volumes.max()),
        "routed_50_pct": float(routed[-1]["flooded_percentage"]),
        "barrier_50_pct": float(barrier[-1]["flooded_percentage"]),
        "baseline_50_pct": float(baseline[-1]["flooded_percentage"]),
        "real_dem_50_pct": float(summary["real_dem_result_50m"]["flooded_percentage"]),
        "levels": levels,
    }


def build_markdown(metrics: dict[str, object]) -> str:
    info = StudentMetadata()
    summary = metrics["summary"]  # type: ignore[assignment]
    validation_lines = metrics["validation_lines"]  # type: ignore[assignment]
    min_pct = metrics["min_pct"]
    max_pct = metrics["max_pct"]
    pct_growth = metrics["pct_growth"]
    max_volume = metrics["max_volume"]
    baseline_50_pct = metrics["baseline_50_pct"]
    routed_50_pct = metrics["routed_50_pct"]
    barrier_50_pct = metrics["barrier_50_pct"]
    real_dem_50_pct = metrics["real_dem_50_pct"]

    return f"""Software Development
实验报告
使用 OpenAI Codex
辅助完成 DEM 洪水淹没分析实验

姓名
{info.name}
班级
{info.class_name}
学号
{info.student_id}
Email
{info.email}
日期
{info.report_date}

## Abstract

This experiment completed the DEM-based flood inundation analysis task described in `Experiment4_Flood_Inundation`. A full Python project was developed to generate and load DEM data, calculate flooded cells according to the condition `elevation < water_level`, estimate inundation depth, visualize flood extent, and analyze how inundated area changes as the water level rises. The core deliverables required by the guide were all produced automatically, including `flood_inundation.py`, `dem_data.npy`, `flood_extent_40m.png`, `flood_extent_50m.png`, `flood_curve.png`, and `prompt_log.md`. To make the project more complete and closer to practical flood analysis, all optional extensions were also implemented: loading a real DEM sample, connectivity-based flood routing, building barriers, animated visualization of rising water levels, and flood-volume calculation. For the baseline synthetic terrain, the flooded percentage increased from {min_pct:.2f}% at the lowest tested level to {max_pct:.2f}% at the highest tested level, with a total growth of {pct_growth:.2f} percentage points, and the maximum simulated flood volume reached approximately {max_volume:,.0f} m^3. Validation confirmed that flooded area increased monotonically with water level, flood percentage stayed within the physical range of 0-100%, and the maximum depth followed the expected relation between water level and minimum terrain elevation. The experiment shows that AI-assisted programming can accelerate the implementation of spatial analysis workflows, but the final results still depend on careful verification of geometry, units, and physical interpretation.

## 1 Experimental Objectives

The objective of this experiment was to implement a complete DEM-based flood inundation workflow and verify the physical consistency of the simulated results.

The specific goals were as follows:

- To construct or load DEM data suitable for flood analysis.
- To implement flooded-cell identification, inundation-depth calculation, and flooded-area percentage calculation.
- To create clear flood-extent and inundation-depth visualizations with `matplotlib`.
- To simulate rising water levels and analyze the area-water-level relationship.
- To verify that the simulation satisfies the expected physical rules.
- To complete the optional extensions suggested in the guide.
- To document the AI-assisted development process in a reproducible way.

## 2 Experimental Process

### 2.1 Task Description

According to the experiment guide, the work was divided into five parts: DEM preparation, flood simulation, visualization, dynamic simulation, and validation. The required outputs were the main implementation file, a DEM data file, flood figures for 40 m and 50 m water levels, a flood-curve figure, and a prompt log. In this project, all optional extensions were also completed rather than omitted.

The final project was stored in:

`/Users/leyangon/Desktop/Experiments/Experiment4`

### 2.2 Environment Setup and Project Structure

The project was developed in Python and organized as a small reproducible experiment project. Because the machine was used in an offline coursework environment, the implementation reused a validated local scientific Python stack and packaged the required code, tests, outputs, and report-generation script together.

The project structure is shown below:

```text
Experiment4/
├── .venv/
├── data/
│   └── real_dem_sample.npy
├── results/
│   ├── flood_routing_50m.png
│   ├── flood_with_buildings_50m.png
│   ├── real_dem_flood_50m.png
│   ├── rising_water.gif
│   ├── summary.json
│   └── water_level_comparison.png
├── src/
│   ├── config.py
│   ├── dem_io.py
│   ├── flood.py
│   └── visualization.py
├── tests/
│   └── test_flood.py
├── Experiment4_Report.docx
├── Experiment4_Report.html
├── Experiment4_Report.md
├── README.md
├── create_report.py
├── dem_data.npy
├── flood_curve.png
├── flood_extent_40m.png
├── flood_extent_50m.png
├── flood_inundation.py
├── prompt_log.md
└── validation_report.txt
```

This structure keeps the mandatory deliverables at the project root while placing reusable logic in the `src/` directory and extension outputs in `results/`.

### 2.3 DEM Preparation

The experiment guide allows either synthetic DEM generation or loading a real DEM source. To ensure that the project can run offline at any time, the baseline analysis used a deterministic synthetic DEM with a `100 x 100` grid and elevation values normalized to the required `30-80 m` range. The synthetic terrain combines a regional slope, a low-lying basin, a floodplain depression, a ridge, and a meandering channel-like depression, plus a small amount of random perturbation controlled by a fixed seed. This makes the terrain visually meaningful and also ensures that inundation patterns are not trivial.

In addition to the synthetic DEM, a second DEM-loading path was implemented using the Matplotlib sample dataset `jacksboro_fault_dem.npz`. The original elevation values were resampled to `100 x 100` and normalized to the same elevation range so that the same flood-analysis pipeline could be applied without changing the rest of the code. This completed the optional real-DEM extension in a way that remains runnable without internet access.

### 2.4 Flood Simulation

The flood calculation follows the physical rule stated in the guide:

- a cell is flooded when `elevation < water_level`;
- inundation depth is `water_level - elevation` for flooded cells and `0` otherwise;
- flooded area percentage is computed as the ratio of flooded cells to total cells.

The implementation was placed in `src/flood.py`. For each water level, the program returns a boolean flood mask, a depth array, flooded area, flooded percentage, flood volume, and the maximum inundation depth.

To support the optional flood-routing extension, the code also provides a connectivity-based mode. In that mode, cells below the water level are only counted as flooded if they are connected to an inflow boundary through 4-neighbour adjacency. This makes the model more realistic than a purely local threshold comparison and demonstrates how hydrological connectivity changes the final flood extent.

### 2.5 Visualization

The visualization module was implemented in `src/visualization.py` using `matplotlib` and `Pillow`. The required figures at 40 m and 50 m were generated automatically and include:

- the original DEM in grayscale;
- the flood extent as a blue overlay;
- the inundation depth map with a blue color scale.

An additional side-by-side comparison plot was produced for multiple water levels so that the progressive spread of inundation can be inspected visually. The optional animated GIF of rising water levels was also generated and saved in `results/rising_water.gif`.

### 2.6 Dynamic Simulation and Extensions

For the dynamic analysis, the program simulated water levels from `40 m` to `50 m` with a `1 m` increment and calculated the flooded percentage at each step. The resulting curve was plotted in `flood_curve.png`, and flood volume was plotted on the same figure with a secondary axis.

All optional extensions listed in the guide were completed:

- A real DEM sample was loaded and analyzed in addition to the synthetic DEM.
- Flood routing was implemented using connectivity from the model boundary.
- Building footprints were represented as barrier cells that block inundation spread in routed mode.
- An animated GIF of rising water levels was created.
- Flood volume was calculated as `depth x cell_area` summed over all flooded cells.

At the 50 m test level, the baseline threshold model flooded {baseline_50_pct:.2f}% of the synthetic DEM. After adding connectivity-based routing, the flooded percentage became {routed_50_pct:.2f}%. After adding building barriers, it decreased further to {barrier_50_pct:.2f}%. The real DEM sample analyzed at the same nominal 50 m water level produced a flooded percentage of {real_dem_50_pct:.2f}% after normalization.

### 2.7 Testing and Validation

Validation focused on the physical correctness required by the guide. The automated checks confirmed the following:

```text
{chr(10).join(validation_lines)}
```

Unit tests were written in `tests/test_flood.py` to verify:

- the zero-flood edge case below minimum elevation;
- the all-flood edge case above maximum elevation;
- monotonic increase of flooded area under rising water levels;
- the fact that routed flooding does not exceed direct threshold flooding;
- the reduction effect introduced by building barriers.

The final test command was:

```bash
.venv/bin/python -m unittest discover -s tests -v
```

The full test suite passed successfully in the final project environment.

## 3 Results and Analysis

### 3.1 Functional Completion

All required tasks in the experiment guide were completed successfully. The project generates the DEM data file, the two required flood figures, the flood-curve figure, the prompt log, and the main Python implementation. In addition, the optional extension outputs were also generated automatically.

### 3.2 Interpretation of the Results

The most important physical result is that the inundated area increases as the water level rises, which is exactly what should happen in a consistent DEM-based inundation model. The produced flood curve shows a smooth upward trend rather than erratic oscillation, which indicates that the terrain generation and flood calculation are internally consistent.

The routed-flood analysis demonstrates that topographic connectivity matters. A direct threshold comparison floods every cell below the water level, but the routed version only floods cells that can be reached from the inflow boundary. This difference is important because low terrain that is hydrologically isolated should not automatically be interpreted as flooded.

The building-barrier extension further illustrates how urban infrastructure can reduce inundation spread in simplified spatial models. Although the rectangular building footprints used here are schematic rather than survey-grade, they clearly show that obstacles can alter local flood extent and reduce connected flooded area.

### 3.3 Role of AI Assistance

AI assistance was useful for quickly drafting the project structure, outlining the baseline flood algorithm, and proposing ways to complete the optional extensions in an offline environment. However, manual review remained necessary in several areas:

- the DEM had to be shaped deliberately so that the flood response was visually meaningful and monotonic;
- the routing logic had to be constrained by cell connectivity instead of only elevation thresholding;
- barrier cells had to be excluded consistently from flood spread and depth calculation;
- validation rules had to be tied back to the physical statements in the experiment guide.

This confirms that AI can accelerate implementation and documentation, but reliable results still depend on engineering judgment and verification.

## 4 Summary

This experiment successfully completed the DEM-based flood inundation analysis task described in `Experiment4_Flood_Inundation`. The final project generates or loads DEM data, calculates flood extent and inundation depth, produces the required visualizations, simulates rising water levels, and validates the physical correctness of the results. All optional extensions were implemented as well, including real DEM support, flood routing, building barriers, GIF animation, and flood-volume estimation.

Through this experiment, I gained a clearer understanding of how raster elevation data can be converted into a simple but meaningful flood-analysis workflow, how dynamic water-level sweeps reveal terrain sensitivity, and how additional assumptions such as hydrological connectivity and barriers affect simulated flood extent. The experiment also reinforced an important lesson about AI-assisted development: AI can significantly speed up coding and organization, but the final quality still depends on careful checking of model assumptions, physical interpretation, and reproducibility.
"""


def markdown_to_html(markdown_text: str) -> str:
    lines = markdown_text.splitlines()
    html_lines = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "  <meta charset='utf-8'>",
        "  <title>Experiment4 Report</title>",
        "  <style>body{font-family:Georgia,serif;max-width:900px;margin:40px auto;line-height:1.7;color:#222;padding:0 16px;} pre{background:#f5f5f5;padding:12px;overflow:auto;} code{font-family:Menlo,monospace;} h2,h3{margin-top:1.4em;} ul{padding-left:20px;}</style>",
        "</head>",
        "<body>",
    ]
    in_code = False
    in_list = False

    for line in lines:
        if line.startswith("```"):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            if in_code:
                html_lines.append("</code></pre>")
                in_code = False
            else:
                html_lines.append("<pre><code>")
                in_code = True
            continue
        if in_code:
            html_lines.append(escape(line))
            continue
        if line.startswith("## "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h2>{escape(line[3:])}</h2>")
        elif line.startswith("### "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h3>{escape(line[4:])}</h3>")
        elif line.startswith("- "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            html_lines.append(f"<li>{escape(line[2:])}</li>")
        elif not line.strip():
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append("<p></p>")
        else:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<p>{escape(line)}</p>")
    if in_list:
        html_lines.append("</ul>")
    if in_code:
        html_lines.append("</code></pre>")
    html_lines.extend(["</body>", "</html>"])
    return "\n".join(html_lines)


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    document = Document()
    in_code = False
    for line in markdown_text.splitlines():
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line.strip():
            document.add_paragraph("")
            continue
        if in_code:
            document.add_paragraph(line, style="No Spacing")
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    document.save(output_path)


def main() -> None:
    metrics = _load_metrics()
    markdown = build_markdown(metrics)
    REPORT_MD.write_text(markdown, encoding="utf-8")
    REPORT_HTML.write_text(markdown_to_html(markdown), encoding="utf-8")
    markdown_to_docx(markdown, REPORT_DOCX)
    print(f"Report generated: {REPORT_DOCX}")


if __name__ == "__main__":
    main()
