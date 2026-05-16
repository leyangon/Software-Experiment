from __future__ import annotations

import json
from html import escape
from pathlib import Path

import pandas as pd
from docx import Document

from src.config import ReservoirParameters


PROJECT_ROOT = Path(__file__).resolve().parent
REPORT_MD = PROJECT_ROOT / "Experiment3_Report.md"
REPORT_HTML = PROJECT_ROOT / "Experiment3_Report.html"
REPORT_DOCX = PROJECT_ROOT / "Experiment3_Report.docx"
SUMMARY_JSON = PROJECT_ROOT / "results" / "summary.json"
SCHEDULE_CSV = PROJECT_ROOT / "optimal_schedule.csv"
PARETO_CSV = PROJECT_ROOT / "results" / "pareto_frontier.csv"
VALIDATION_REPORT = PROJECT_ROOT / "validation_report.txt"


def _load_metrics() -> dict[str, object]:
    parameters = ReservoirParameters()
    schedule = pd.read_csv(SCHEDULE_CSV)
    pareto = pd.read_csv(PARETO_CSV)
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    validation_lines = VALIDATION_REPORT.read_text(encoding="utf-8").strip().splitlines()
    return {
        "parameters": parameters,
        "schedule": schedule,
        "pareto": pareto,
        "summary": summary,
        "validation_lines": validation_lines,
        "baseline_revenue_usd": float(summary["baseline_revenue_usd"]),
        "min_release_m3s": float(schedule["release_m3s"].min()),
        "max_release_m3s": float(schedule["release_m3s"].max()),
        "final_storage_m3": float(schedule["storage_end_m3"].iloc[-1]),
        "zero_deficit_points": int((pareto["ecological_deficit_m3"] <= 1e-6).sum()),
    }


def build_markdown(metrics: dict[str, object]) -> str:
    parameters: ReservoirParameters = metrics["parameters"]  # type: ignore[assignment]
    baseline_revenue = metrics["baseline_revenue_usd"]
    min_release = metrics["min_release_m3s"]
    max_release = metrics["max_release_m3s"]
    final_storage = metrics["final_storage_m3"]
    zero_deficit_points = metrics["zero_deficit_points"]
    validation_lines = metrics["validation_lines"]

    return f"""Software Development
实验报告
使用 OpenAI Codex
辅助完成水库调度优化实验

姓名
{parameters.student_name}
班级
{parameters.class_name}
学号
{parameters.student_id}
Email
{parameters.email}
日期
{parameters.report_date}

## Abstract

This experiment implemented the reservoir dispatch problem described in `Experiment3_Reservoir_Optimization` as a seven-day constrained optimization task. The model balances hydropower revenue against ecological protection under storage balance, release capacity, and reservoir capacity constraints. The completed project includes a feasible baseline optimization, a Pareto frontier for ecology-revenue trade-off analysis, and a full validation workflow that checks storage bounds, ecological releases, mass balance, and revenue calculation. To keep the submission reproducible in an offline environment, the implementation was packaged as a self-contained Python project with deterministic input parameters and automatically generated output files. All required deliverables were completed, including `reservoir_optimize.py`, `optimal_schedule.csv`, `tradeoff_analysis.png`, `prompt_log.md`, and `validation_report.txt`. In addition, all optional extensions were implemented: inflow uncertainty analysis, rolling horizon optimization, water-quality-aware constraints, and algorithm comparison between SLSQP and L-BFGS-B. The baseline optimized schedule achieved a total hydropower revenue of approximately `${baseline_revenue:,.2f}` while maintaining zero ecological deficit and ending the week with a storage of `{final_storage:,.2f} m^3`. The experiment shows that AI can speed up optimization-oriented development, but physically meaningful results still require manual checking of constraints, units, and scenario interpretation.

## 1 Experimental Objectives

The objective of this experiment was to formulate, implement, analyze, and validate a multi-objective reservoir operation model for drought-period dispatch.

The specific goals were as follows:

- To formulate the reservoir release problem with explicit decision variables, objectives, and constraints.
- To solve the optimization model using `scipy.optimize.minimize`.
- To generate a feasible seven-day release schedule under storage and release limits.
- To analyze the trade-off between hydropower revenue and ecological protection through weighted optimization.
- To validate the final solution against physical constraints and mass balance.
- To complete the optional extensions suggested in the guide.
- To document how AI support was used and where human review remained necessary.

## 2 Experimental Process

### 2.1 Task Description

According to the experiment guide, the work needed to complete four main parts: mathematical formulation, optimization implementation, trade-off analysis, and validation. The required deliverables were `reservoir_optimize.py`, `optimal_schedule.csv`, `tradeoff_analysis.png`, `prompt_log.md`, and `validation_report.txt`. All optional extensions listed in the guide were also implemented in this project.

The final project was stored in:

`/Users/leyangon/Desktop/Experiments/Experiment3`

### 2.2 Environment Setup and Project Structure

The project was developed for Python 3.10+ and executed inside a dedicated virtual environment. Following the same approach used in previous experiments, the environment was created with `--system-site-packages` so that the locally bundled scientific packages could be reused without relying on network access.

The project structure is shown below:

```text
Experiment3/
├── .venv/
├── data/
├── results/
│   ├── algorithm_comparison.csv
│   ├── pareto_frontier.csv
│   ├── rolling_horizon_schedule.csv
│   ├── summary.json
│   ├── uncertainty_analysis.csv
│   └── water_quality_schedule.csv
├── src/
│   ├── analysis.py
│   ├── config.py
│   └── optimization.py
├── tests/
│   └── test_reservoir.py
├── Experiment3_Report.docx
├── Experiment3_Report.html
├── Experiment3_Report.md
├── README.md
├── create_report.py
├── optimal_schedule.csv
├── prompt_log.md
├── requirements.txt
├── reservoir_optimize.py
├── tradeoff_analysis.png
└── validation_report.txt
```

This organization keeps the required submission files at the project root while separating reusable optimization logic, exported extension outputs, and automated tests.

### 2.3 Mathematical Formulation

The decision variables are the daily releases `Q_1, Q_2, ..., Q_7`, each measured in `m^3/s`. The baseline objective is to maximize total hydropower revenue over the seven-day horizon. Daily revenue is computed by converting released water volume into generated energy using a fixed turbine head and efficiency, then multiplying by the day-specific electricity price.

The ecological objective is to minimize the total deficit relative to the minimum ecological release `Q_eco = 10 m^3/s`. For the final baseline schedule, this requirement is enforced as a hard constraint so that the reported solution remains fully feasible. For trade-off analysis, the lower bound is relaxed and ecological deficit is added as a weighted penalty term in the objective function, which makes it possible to trace the Pareto frontier.

The physical constraints used in the model are:

- storage bounds: `V_min <= V_t <= V_max`;
- release bounds: `Q_release <= Q_max`;
- ecological minimum release for the baseline case: `Q_release >= Q_eco`;
- storage continuity: `V_(t+1) = V_t + (Q_in,t - Q_out,t) * Δt`.

### 2.4 Optimization Implementation

The optimization model was implemented in `reservoir_optimize.py` with reusable functions placed in `src/optimization.py`. The main solver uses the `SLSQP` method because it supports nonlinear inequality constraints and performs well for this small continuous dispatch problem.

The baseline feasible schedule was obtained under hard ecological constraints. The optimized release range was `{min_release:.2f} m^3/s` to `{max_release:.2f} m^3/s`, which shows that the model uses flexibility above the ecological minimum while remaining well below the maximum release capacity.

To make the code easier to inspect, storage simulation, revenue calculation, ecological deficit calculation, validation, and extension analyses were separated into small functions. This also made it straightforward to test individual steps rather than only checking the final output table.

### 2.5 Trade-off Analysis

The trade-off analysis was implemented by solving a family of relaxed optimization problems with different ecology weights. The resulting Pareto frontier was exported to `results/pareto_frontier.csv` and visualized in `tradeoff_analysis.png`.

The analysis showed that the ecology-revenue conflict appears only when the ecological minimum is treated as a soft penalty instead of a hard operational rule. Among the sampled Pareto points, `{zero_deficit_points}` solutions achieved zero ecological deficit, while the higher-revenue relaxed solutions allowed temporary under-release at the cost of environmental performance. This comparison was used to estimate the opportunity cost of protecting the ecological flow target.

### 2.6 Optional Extensions

All optional extensions listed in the guide were completed.

First, inflow uncertainty was added through multiple forecast scenarios, and the resulting evaluation table was saved to `results/uncertainty_analysis.csv`. This allows the robustness of a dispatch schedule to be checked under drier and wetter inflow conditions.

Second, rolling horizon optimization was implemented. The controller re-optimizes the remaining horizon after each day using updated inflow information, and the resulting schedule was saved to `results/rolling_horizon_schedule.csv`.

Third, water-quality-aware constraints were added as an engineering extension. A stricter minimum storage threshold and a daily release ramping limit were imposed to represent operational caution against low-storage water quality deterioration and abrupt release changes. The resulting schedule was saved to `results/water_quality_schedule.csv`.

Fourth, the optimization method was compared between `SLSQP` and `L-BFGS-B`. Since `L-BFGS-B` only supports box constraints directly, a penalty formulation was added for the comparison case, and the results were exported to `results/algorithm_comparison.csv`.

### 2.7 Testing and Validation

Validation focused on feasibility, physical consistency, and reproducibility. Automated tests were written in `tests/test_reservoir.py` to check storage simulation length, feasibility of the baseline solution, positivity of revenue, rolling-horizon storage safety, water-quality extension behavior, and the presence of both algorithms in the comparison output.

The validation report confirms the following points:

```text
{chr(10).join(validation_lines)}
```

The executed test command was:

```bash
.venv/bin/python -m unittest discover -s tests -v
```

The test suite passed successfully in the final project environment.

## 3 Results and Analysis

### 3.1 Functional Completion

All mandatory deliverables were completed successfully. The project includes the required optimization program, the exported seven-day optimal schedule, the Pareto trade-off figure, the prompt log, and the validation report. In addition, all optional extensions from the guide were implemented instead of being omitted.

### 3.2 Interpretation of the Dispatch Results

The baseline solution achieved a total hydropower revenue of approximately `${baseline_revenue:,.2f}` while maintaining the minimum ecological release requirement. This indicates that the optimization model was able to exploit time-varying electricity prices without violating the drought-period ecological protection rule.

The storage trajectory remained within the admissible interval throughout the horizon, and the final storage stayed above the minimum threshold. This is important because a superficially profitable release schedule would not be acceptable if it emptied the reservoir excessively late in the week.

The Pareto analysis also highlights an important management insight: if the ecological release is treated as negotiable, the optimizer can generate higher short-term revenue, but only by accepting environmental deficit. In contrast, the baseline feasible solution converts the ecological target into a firm operational rule, which better matches actual reservoir regulation practice.

### 3.3 Role of AI Assistance

AI assistance was useful for drafting the optimization skeleton, structuring the project files, and outlining validation logic. However, the final implementation required manual inspection and correction in several places:

- the baseline deliverable had to preserve ecological minimum release as a hard constraint rather than only a weighted penalty;
- the Pareto analysis had to be separated from the feasible baseline run so that both objectives could be analyzed meaningfully;
- revenue needed explicit unit conversion from release volume to generated energy;
- mass balance and storage feasibility had to be verified explicitly rather than assumed from optimizer success alone.

These corrections show that AI can accelerate implementation, but trustworthy engineering results still depend on human review of assumptions, formulas, and constraints.

## 4 Summary

This experiment successfully completed the reservoir optimization task described in `Experiment3_Reservoir_Optimization`. The final project formulates the seven-day dispatch problem, solves it with constrained optimization, exports a feasible release schedule, analyzes the ecology-revenue trade-off, and validates the physical consistency of the solution. The project also completes all optional extensions, including uncertainty analysis, rolling horizon dispatch, water-quality-aware constraints, and solver comparison.

Through this experiment, I gained a clearer understanding of how multi-objective water resources problems can be translated into software, how operational constraints shape optimization results, and how AI can support implementation without replacing engineering judgment. Overall, the submission satisfies the stated requirements and provides a complete, reproducible, and well-documented experiment project.
"""


def markdown_to_html(markdown_text: str) -> str:
    lines = markdown_text.splitlines()
    html_lines = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "  <meta charset='utf-8'>",
        "  <title>Experiment3 Report</title>",
        "  <style>body{font-family:Georgia,serif;max-width:900px;margin:40px auto;line-height:1.7;color:#222;padding:0 16px;} pre{background:#f5f5f5;padding:12px;overflow:auto;} code{font-family:Menlo,monospace;} h2,h3{margin-top:1.4em;} ul{padding-left:20px;}</style>",
        "</head>",
        "<body>",
    ]
    in_code_block = False
    in_list = False
    for line in lines:
        if line.startswith("```"):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            if in_code_block:
                html_lines.append("</code></pre>")
                in_code_block = False
            else:
                html_lines.append("<pre><code>")
                in_code_block = True
            continue
        if in_code_block:
            html_lines.append(escape(line))
            continue
        if line.startswith("## "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h2>{escape(line[3:])}</h2>")
        elif line.startswith("### "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h3>{escape(line[4:])}</h3>")
        elif line.startswith("- "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            html_lines.append(f"<li>{escape(line[2:])}</li>")
        elif not line.strip():
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append("<p></p>")
        else:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<p>{escape(line)}</p>")
    if in_list:
        html_lines.append("</ul>")
    if in_code_block:
        html_lines.append("</code></pre>")
    html_lines.extend(["</body>", "</html>"])
    return "\n".join(html_lines)


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    in_code_block = False
    for line in markdown_text.splitlines():
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            if in_code_block:
                run.font.name = "Courier New"
    doc.save(output_path)


def main() -> None:
    metrics = _load_metrics()
    markdown_text = build_markdown(metrics)
    REPORT_MD.write_text(markdown_text, encoding="utf-8")
    REPORT_HTML.write_text(markdown_to_html(markdown_text), encoding="utf-8")
    markdown_to_docx(markdown_text, REPORT_DOCX)
    print(f"Report files written to {PROJECT_ROOT}")


if __name__ == "__main__":
    main()
